"""
Génère le rapport de projet VaR — M2 Finance EFREI
Auteurs : Sidy Laye Sarr, Moulaye Koutam, Abdoulaye Diop
"""
import sys, os, math
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Helpers ───────────────────────────────────────────────────────────────────
def rgb(r,g,b): return RGBColor(r,g,b)

BLUE   = rgb(31,  78, 121)
RED    = rgb(192,   0,   0)
GREEN  = rgb(55,  86,  35)
GRAY   = rgb(89,  89,  89)
LGRAY  = rgb(166,166,166)

def add_heading(doc, text, level=1, color=BLUE, space_before=18, space_after=6):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color
    run.font.bold = True
    run.font.size = Pt([0, 18, 15, 13, 12][level])
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def add_para(doc, text, bold=False, italic=False, color=None,
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_formula(doc, formula, explanation=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(formula)
    r.font.name  = "Courier New"
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = BLUE
    if explanation:
        r2 = p.add_run(f"     ({explanation})")
        r2.font.size = Pt(10)
        r2.font.color.rgb = GRAY
    return p

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table_row(table, cells, bold=False, bg=None, center=False):
    row = table.add_row()
    for i, txt in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(txt)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0] if p.runs else p.add_run(str(txt))
        run.bold      = bold
        run.font.size = Pt(10)
        if bg: shade_cell(cell, bg)
    return row

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def hdr_row(table, headers, widths_cm=None):
    row = table.add_row()
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(h)
        run.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255,255,255)
        shade_cell(cell, "1F4E79")
    if widths_cm:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])
    return row

# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()

# Marges
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── PAGE DE GARDE ─────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EFREI PARIS — M2 Finance & Gestion des Risques")
run.font.size = Pt(12); run.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RAPPORT DE PROJET")
run.font.size = Pt(14); run.bold = True; run.font.color.rgb = GRAY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Calcul et Analyse de la\nValue at Risk Historique")
run.font.size = Pt(28); run.bold = True; run.font.color.rgb = BLUE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Conception d'une plateforme web de gestion du risque de portefeuille")
run.font.size = Pt(13); run.italic = True; run.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Auteurs
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sidy Laye Sarr  ·  Moulaye Koutam  ·  Abdoulaye Diop")
run.font.size = Pt(14); run.bold = True; run.font.color.rgb = BLUE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Année universitaire 2025 – 2026")
run.font.size = Pt(11); run.font.color.rgb = GRAY

doc.add_page_break()

# ── RÉSUMÉ EXÉCUTIF ───────────────────────────────────────────────────────────
add_heading(doc, "Résumé exécutif", 1)
add_para(doc,
    "Ce projet présente la conception et le développement d'une plateforme web complète de calcul "
    "de la Value at Risk (VaR) historique 1 jour pour des portefeuilles d'actions cotées. "
    "L'application permet à un investisseur de constituer un portefeuille multi-actifs (jusqu'à "
    "20 titres issus du CAC 40, S&P 500, EuroStoxx 50 et SBF 120), de définir ses pondérations "
    "— soit manuellement, soit par valorisation au cours actuel — et d'obtenir instantanément les "
    "métriques de risque conformes à la méthodologie enseignée en cours."
)
add_para(doc,
    "La méthodologie implémentée suit strictement la méthode historique du cours Stat M2 EFREI : "
    "tri des rendements journaliers, sélection du k-ième pire rendement par la fonction SMALL "
    "(k = ⌊n × q⌋), et calcul du CVaR comme moyenne des k−1 pires rendements. "
    "La fenêtre de calcul de la VaR est fixée à 20 ans glissants pour garantir la robustesse "
    "statistique, tandis que la fenêtre de performance est définie librement par l'utilisateur."
)
add_para(doc,
    "L'application est développée avec FastAPI (Python) pour le backend et React/TypeScript "
    "pour le frontend, et consomme les données historiques en temps réel via l'API Yahoo Finance."
)

# ── RÉPARTITION DES CONTRIBUTIONS ────────────────────────────────────────────
add_heading(doc, "1. Répartition des contributions", 1)
add_para(doc,
    "Le projet a été réalisé en trinôme avec une répartition équitable du travail. "
    "Chaque membre a pris en charge un tiers du projet, couvrant à la fois des aspects "
    "théoriques, techniques et analytiques."
)

t_contrib = doc.add_table(rows=1, cols=3)
t_contrib.style = "Table Grid"
t_contrib.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(t_contrib, ["Auteur", "Domaine", "Sections couvertes"], [4, 5, 9])

contrib_data = [
    (
        "Sidy Laye Sarr",
        "Méthodologie & Théorie",
        (
            "• Section 3 — Méthodologie VaR historique complète : fondements théoriques, "
            "formules (SMALL, FLOOR, quantile), double fenêtre temporelle\n"
            "• Implémentation du module var_calculator.py (calcul k, tri des rendements, "
            "filtrage des outliers, corrélations)\n"
            "• Rédaction de la documentation mathématique et des formules du rapport"
        ),
    ),
    (
        "Moulaye Koutam",
        "Architecture & Backend",
        (
            "• Section 4 — Architecture de la plateforme : stack technique, API FastAPI, "
            "modèles de données SQLAlchemy, endpoints REST\n"
            "• Implémentation du data_fetcher.py (Yahoo Finance, curl_cffi, gestion du crumb)\n"
            "• Mise en place de la base de données, du cache des métriques et des routers "
            "(var.py, portfolios.py, stocks.py)\n"
            "• Section 5 — Fonctionnalités : sélection des actifs, valorisation par nombre "
            "d'actions, dashboard analytique"
        ),
    ),
    (
        "Abdoulaye Diop",
        "Frontend & Analyse",
        (
            "• Section 6 — Cas d'application : analyse complète AI.PA + AIR.PA "
            "(VaR = 2 343 €, rendement +96,09 %), interprétation des résultats\n"
            "• Développement du frontend React/TypeScript : stepper 4 étapes, "
            "composants VarCard, ReturnHistogram, PerformanceChart, Dashboard\n"
            "• Section 7 — Export Excel avec formules natives (SMALL, INDEX, AVERAGE, FLOOR)\n"
            "• Section 8 — Conclusion et perspectives"
        ),
    ),
]

for name, domain, details in contrib_data:
    row = t_contrib.add_row()
    # Colonne Auteur
    cell_name = row.cells[0]
    cell_name.text = name
    p = cell_name.paragraphs[0]
    run = p.runs[0] if p.runs else p.add_run(name)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    # Colonne Domaine
    cell_domain = row.cells[1]
    cell_domain.text = domain
    p2 = cell_domain.paragraphs[0]
    run2 = p2.runs[0] if p2.runs else p2.add_run(domain)
    run2.bold = True
    run2.font.size = Pt(10)
    # Colonne Détails
    cell_det = row.cells[2]
    cell_det.text = details
    p3 = cell_det.paragraphs[0]
    run3 = p3.runs[0] if p3.runs else p3.add_run(details)
    run3.font.size = Pt(9)

doc.add_paragraph()

# ── TABLE DES MATIÈRES ────────────────────────────────────────────────────────
add_heading(doc, "Table des matières", 1)
toc_items = [
    ("1.", "Répartition des contributions"),
    ("2.", "Contexte et objectifs"),
    ("3.", "Méthodologie VaR Historique"),
    ("   3.1", "Définition et fondements théoriques"),
    ("   3.2", "Calcul des rendements"),
    ("   3.3", "VaR historique : méthode SMALL"),
    ("   3.4", "CVaR / Expected Shortfall"),
    ("   3.5", "Double fenêtre temporelle"),
    ("4.", "Architecture de la plateforme"),
    ("   4.1", "Stack technique"),
    ("   4.2", "Architecture backend (FastAPI)"),
    ("   4.3", "Architecture frontend (React)"),
    ("5.", "Fonctionnalités de l'application"),
    ("6.", "Cas d'application : AI.PA + AIR.PA"),
    ("7.", "Export Excel et formules"),
    ("8.", "Conclusion"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{num}  ")
    r1.font.bold  = ("." in num and len(num) <= 3)
    r1.font.size  = Pt(11)
    r1.font.color.rgb = BLUE if ("." in num and len(num) <= 3) else GRAY
    r2 = p.add_run(title)
    r2.font.size  = Pt(11)
    r2.font.bold  = ("." in num and len(num) <= 3)
    r2.font.color.rgb = BLUE if ("." in num and len(num) <= 3) else GRAY

doc.add_page_break()

# ── 1. CONTEXTE ET OBJECTIFS ─────────────────────────────────────────────────
add_heading(doc, "2. Contexte et objectifs", 1)
add_para(doc,
    "La gestion du risque de marché est au cœur des exigences réglementaires issues de Bâle II et "
    "Bâle III. La Value at Risk (VaR) constitue l'indicateur standard de mesure du risque de perte "
    "maximale d'un portefeuille sur un horizon donné, à un niveau de confiance fixé. Dans le cadre "
    "de notre formation M2 Finance à l'EFREI, nous avons développé une application web permettant "
    "de calculer, visualiser et sauvegarder des analyses VaR historiques pour des portefeuilles "
    "d'actions."
)

add_heading(doc, "Objectifs du projet", 2, color=BLUE)
for obj in [
    "Implémenter la méthode VaR historique conforme aux formules du cours (méthode SMALL, k = ⌊n × q⌋)",
    "Développer une interface utilisateur intuitive guidant l'investisseur étape par étape",
    "Intégrer les données de marché en temps réel via Yahoo Finance (CAC 40, S&P 500, EuroStoxx 50, SBF 120)",
    "Permettre la valorisation d'un portefeuille par nombre d'actions × cours actuel",
    "Exporter les calculs vers Excel avec les formules natives (SMALL, AVERAGE, INDEX)",
    "Sauvegarder et comparer les analyses de risque sur un tableau de bord analytique",
]:
    add_bullet(doc, obj)

# ── 2. MÉTHODOLOGIE ───────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "3. Méthodologie VaR Historique", 1)

add_heading(doc, "3.1  Définition et fondements théoriques", 2, color=BLUE)
add_para(doc,
    "La Value at Risk historique 1 jour répond à la question : "
    "« Quelle est la perte maximale que mon portefeuille peut subir en une journée, "
    "avec une probabilité de (1 − q) ? »"
)
add_para(doc,
    "Formellement, pour un portefeuille de valeur V et un quantile q ∈ (0,1) :",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_formula(doc, "VaR_q  =  − inf { r : P(R ≤ r) ≥ q }",
            "q = 0,05 pour un niveau de confiance de 95%")
add_para(doc,
    "La méthode historique n'impose aucune hypothèse paramétrique sur la distribution des "
    "rendements : elle utilise directement les données de marché observées. C'est la méthode "
    "retenue par les régulateurs (Bâle III) et celle enseignée dans le cours Stat M2 EFREI."
)

add_heading(doc, "3.2  Calcul des rendements journaliers", 2, color=BLUE)
add_para(doc,
    "Pour chaque action i, le rendement journalier est calculé comme la variation relative "
    "du cours de clôture ajusté (dividendes et splits inclus) entre deux jours consécutifs :"
)
add_formula(doc, "r_{i,t}  =  (P_{i,t} − P_{i,t-1}) / P_{i,t-1}",
            "variation relative du cours ajusté")
add_para(doc,
    "Un filtre de cohérence est appliqué : les rendements dont la valeur absolue dépasse 40 % "
    "sont exclus, car ils correspondent à des erreurs de données (reconstitutions pré-IPO, "
    "données synthétiques antérieures à la cotation). Ce seuil garantit l'intégrité statistique "
    "sans éliminer aucun choc de marché réel (krach de 2008, COVID-19)."
)

add_para(doc, "Le rendement journalier du portefeuille est ensuite calculé comme la somme pondérée :")
add_formula(doc, "R_t  =  Σᵢ  wᵢ × r_{i,t}",
            "wᵢ = poids de l'action i, Σwᵢ = 1")

add_heading(doc, "3.3  VaR historique : méthode SMALL", 2, color=BLUE)
add_para(doc,
    "La méthode implémentée suit strictement la formule du cours, identique à la fonction "
    "Excel SMALL. On trie les n rendements journaliers du portefeuille par ordre croissant "
    "(du plus négatif au plus positif), puis on sélectionne le k-ième :"
)
add_formula(doc, "k  =  ⌊ n × q ⌋",
            "⌊·⌋ = partie entière inférieure (FLOOR en Excel)")
add_formula(doc, "VaR  =  R_trié[k]",
            "k-ième pire rendement = SMALL(R, k)")
add_formula(doc, "VaR_€  =  |VaR| × Montant_portefeuille",
            "perte exprimée en euros")

add_para(doc,
    "Exemple : sur une fenêtre VaR de 5 115 observations, avec q = 5 % :"
)

# Tableau exemple calcul
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(t, ["Paramètre", "Calcul", "Valeur"], [5, 7, 5])
for row_data in [
    ("n (observations)", "COUNTA(rendements)", "5 115"),
    ("q (quantile)", "choix utilisateur", "5 %"),
    ("k", "⌊5 115 × 0,05⌋", "255"),
    ("VaR 95% 1J (%)", "SMALL(R, 255)", "−2,3430 %"),
    ("VaR 95% 1J (€)", "|−2,3430 %| × 100 000 €", "2 343,00 €"),
]:
    add_table_row(t, row_data)
doc.add_paragraph()

add_heading(doc, "3.4  CVaR / Expected Shortfall", 2, color=BLUE)
add_para(doc,
    "Le CVaR (Conditional Value at Risk), également appelé Expected Shortfall, mesure la perte "
    "moyenne dans les scénarios pires que la VaR. Il est calculé comme la moyenne des "
    "k − 1 rendements les plus défavorables (strictement inférieurs à la VaR) :"
)
add_formula(doc, "CVaR  =  (1/(k−1))  ×  Σⱼ₌₁^{k-1}  R_trié[j]",
            "AVERAGE(k−1 pires rendements) en Excel")
add_formula(doc, "CVaR_€  =  |CVaR| × Montant_portefeuille")
add_para(doc,
    "Cette formule est identique à =AVERAGE(D2:D{k-1}) dans le fichier Excel du cours, "
    "où la colonne D contient les rendements triés par ordre croissant."
)

add_heading(doc, "3.5  Double fenêtre temporelle", 2, color=BLUE)
add_para(doc,
    "L'application distingue deux fenêtres temporelles distinctes, conformément à la "
    "pratique institutionnelle :"
)

t2 = doc.add_table(rows=1, cols=3)
t2.style = "Table Grid"
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(t2, ["Fenêtre", "Période", "Usage"], [4, 7, 6])
for row_data in [
    ("VaR (risque)", "date_fin − 20 ans → date_fin", "Calcul VaR, histogramme des rendements"),
    ("Performance", "choix utilisateur → date_fin", "Rendement total, volatilité, graphique cumulatif"),
]:
    r = t2.add_row()
    for i, txt in enumerate(row_data):
        r.cells[i].text = txt
        r.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()
add_para(doc,
    "La fenêtre VaR de 20 ans garantit un nombre d'observations suffisant (≈ 5 000) pour que "
    "le quantile empirique soit statistiquement robuste, tout en restant dans des conditions "
    "de marché homogènes. La fenêtre de performance laisse à l'utilisateur la liberté d'évaluer "
    "le rendement sur la période de son choix (ex. : depuis l'entrée en portefeuille)."
)

# ── 3. ARCHITECTURE ───────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "4. Architecture de la plateforme", 1)

add_heading(doc, "4.1  Stack technique", 2, color=BLUE)

t3 = doc.add_table(rows=1, cols=3)
t3.style = "Table Grid"
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(t3, ["Couche", "Technologie", "Rôle"], [4, 6, 8])
stack = [
    ("Backend API",       "FastAPI 0.115 (Python 3.11)",  "Endpoints REST : calcul VaR, gestion portefeuilles, prix en temps réel"),
    ("Calcul numérique",  "NumPy 1.26 + Pandas 2.2",     "Tri des rendements, SMALL/AVERAGE, corrélations, statistiques"),
    ("Données marché",    "Yahoo Finance v8 API (curl_cffi)", "Cours ajustés journaliers, recherche de titres mondiaux"),
    ("ORM / BDD",         "SQLAlchemy 2.0 + SQLite",     "Persistance des portefeuilles et cache des métriques VaR"),
    ("Frontend",          "React 18 + TypeScript",        "Interface utilisateur, visualisations, stepper 4 étapes"),
    ("Graphiques",        "Recharts",                     "Histogramme des rendements, courbe de performance"),
    ("Export",            "SheetJS (xlsx) + openpyxl",   "Export Excel frontend (SheetJS) et génération formules (openpyxl)"),
    ("Styles",            "Tailwind CSS",                 "Design system dark mode, composants réutilisables"),
]
for row_data in stack:
    add_table_row(t3, row_data)
doc.add_paragraph()

add_heading(doc, "4.2  Architecture backend", 2, color=BLUE)
add_para(doc,
    "Le backend est organisé en couches distinctes selon le principe de séparation des "
    "responsabilités :"
)
for item in [
    "Routers (app/routers/) : exposition des endpoints REST — /var/calculate, /portfolios/, /stocks/",
    "Services (app/services/) : logique métier — var_calculator.py, data_fetcher.py",
    "Schemas (app/schemas.py) : validation Pydantic des entrées/sorties",
    "Models (app/models.py) : entités SQLAlchemy — Portfolio, PortfolioStock",
    "Database (app/database.py) : session SQLAlchemy asynchrone (aiosqlite)",
]:
    add_bullet(doc, item)

add_para(doc,
    "Le module data_fetcher.py utilise curl_cffi pour simuler un navigateur Chrome et "
    "contourner la détection de bots de Yahoo Finance. En cas d'indisponibilité, un fallback "
    "requests est activé avec récupération du crumb d'authentification. Les données sont "
    "téléchargées depuis le 3 janvier 2000 pour garantir la fenêtre VaR de 20 ans."
)

add_heading(doc, "4.3  Architecture frontend", 2, color=BLUE)
add_para(doc, "L'interface est structurée en 4 pages principales et suit un parcours utilisateur guidé :")

t4 = doc.add_table(rows=1, cols=3)
t4.style = "Table Grid"
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(t4, ["Page", "Composants clés", "Fonctionnalité"], [3, 6, 9])
for row_data in [
    ("Dashboard",    "StatCards, BarChart, Table triable",     "KPIs agrégés, classement des portefeuilles par risque"),
    ("Calculateur",  "StockSelector, WeightManager, ParametersForm", "Stepper 4 étapes : sélection → pondérations → paramètres → résultats"),
    ("Résultats",    "ReturnHistogram, PerformanceChart",      "Histogramme VaR, courbe de valeur cumulée, export Excel"),
    ("Historique",   "PortfolioCard, VarCard",                 "Recalcul, suppression, téléchargement Excel par portefeuille"),
]:
    add_table_row(t4, row_data)
doc.add_paragraph()

# ── 4. FONCTIONNALITÉS ────────────────────────────────────────────────────────
add_heading(doc, "5. Fonctionnalités de l'application", 1)

add_heading(doc, "Sélection des actifs", 2, color=BLUE)
for f in [
    "Recherche instantanée dans la base curatée (CAC 40, S&P 500, EuroStoxx 50, SBF 120)",
    "Recherche en direct sur Yahoo Finance pour tout titre mondial (LVMH, Apple, Siemens, etc.)",
    "Ajout de jusqu'à 20 actions avec pondération égale automatique à l'ajout",
    "Saisie du nombre d'actions par titre et valorisation au cours actuel (bouton « Valoriser »)",
    "Calcul automatique des pondérations à partir des valeurs de marché : wᵢ = (nᵢ × Pᵢ) / Σ(nᵢ × Pᵢ)",
]:
    add_bullet(doc, f)

add_heading(doc, "Calcul et résultats", 2, color=BLUE)
for f in [
    "VaR 95% (ou tout autre quantile) en % et en euros",
    "Histogramme des rendements journaliers avec mise en évidence des queues (VaR en rouge)",
    "Courbe de performance : évolution de la valeur du portefeuille sur la fenêtre choisie",
    "Composition du portefeuille : ticker, nom, poids en %",
    "Sauvegarde avec cache des métriques (PATCH /portfolios/{id}/cache) — pas de re-téléchargement",
]:
    add_bullet(doc, f)

add_heading(doc, "Dashboard analytique", 2, color=BLUE)
for f in [
    "KPIs : nombre de portefeuilles, capital total géré, VaR % moyenne, volatilité moyenne",
    "Graphique en barres horizontales : VaR % par portefeuille, coloré par niveau (Élevé / Modéré / Faible)",
    "Table triable (VaR %, VaR €, volatilité, capital) avec badge de niveau de risque",
    "Classification : Faible (VaR < 1,5%), Modéré (1,5–3%), Élevé (≥ 3%)",
]:
    add_bullet(doc, f)

# ── 5. CAS D'APPLICATION ──────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "6. Cas d'application : Air Liquide + Airbus", 1)

add_para(doc,
    "Afin d'illustrer les capacités de la plateforme, nous avons réalisé une analyse complète "
    "sur un portefeuille équipondéré composé de deux fleurons de la cote française :"
)
add_bullet(doc, "Air Liquide (AI.PA)  — 50 % du portefeuille")
add_bullet(doc, "Airbus Group (AIR.PA) — 50 % du portefeuille")
add_bullet(doc, "Montant investi : 100 000 €")
add_bullet(doc, "Quantile q = 5 % → niveau de confiance 95 %")
add_bullet(doc, "Date de fin : 12 février 2026")
add_bullet(doc, "Fenêtre VaR : 13 février 2006 → 11 février 2026 (20 ans)")
add_bullet(doc, "Fenêtre performance : 4 mars 2021 → 11 février 2026")

add_heading(doc, "Résultats VaR", 2, color=RED)

t5 = doc.add_table(rows=1, cols=3)
t5.style = "Table Grid"
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(t5, ["Indicateur", "Valeur", "Interprétation"], [5, 4, 9])
results = [
    ("n (observations VaR)", "5 115 jours",
     "20 ans de données journalières filtrées"),
    ("k = ⌊5 115 × 0,05⌋", "255",
     "255e pire rendement = VaR"),
    ("VaR 95% 1J (%)", "−2,3430 %",
     "Perte maximale journalière à 95% de confiance"),
    ("VaR 95% 1J (€)", "2 343,00 €",
     "Sur un investissement de 100 000 €"),
    ("CVaR (%)", "−3,6326 %",
     "Perte moyenne si la VaR est dépassée"),
    ("CVaR (€)", "3 632,63 €",
     "Expected Shortfall : moyenne des 254 pires jours"),
]
for row_data in results:
    add_table_row(t5, row_data)
doc.add_paragraph()

add_heading(doc, "Résultats de performance (fenêtre 2021–2026)", 2, color=GREEN)

t6 = doc.add_table(rows=1, cols=2)
t6.style = "Table Grid"
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(t6, ["Indicateur", "Valeur"], [8, 6])
perf = [
    ("Rendement total (5 ans)",     "+96,09 %"),
    ("Valeur finale du portefeuille", "≈ 196 090 €"),
    ("Volatilité annualisée",       "19,83 %"),
    ("Rendement annuel moyen",      "≈ +14,5 %"),
    ("Observations (perf)",         "1 267 jours"),
]
for row_data in perf:
    add_table_row(t6, row_data)
doc.add_paragraph()

add_para(doc,
    "Interprétation : avec un niveau de confiance de 95 %, le portefeuille AI.PA / AIR.PA "
    "ne devrait pas perdre plus de 2 343 € (2,34 % de sa valeur) en une seule journée. "
    "Sur la période 2021–2026, ce portefeuille a plus que doublé de valeur, avec une "
    "volatilité annualisée de 19,83 % — classée « Modérée » selon notre grille d'analyse "
    "(entre 20 % et 30 % → seuil réglementaire Bâle)."
)

add_heading(doc, "Corrélation et diversification", 2, color=BLUE)
add_para(doc,
    "La matrice de corrélation calculée sur la fenêtre performance révèle une corrélation "
    "positive modérée entre les deux titres. Cela confirme l'intérêt de diversifier au-delà "
    "de ces deux valeurs pour réduire davantage le risque de portefeuille. La volatilité "
    "du portefeuille est calculée selon la formule du cours :"
)
add_formula(doc,
    "σₚ = √( w₁²σ₁² + w₂²σ₂² + 2·w₁·w₂·ρ·σ₁·σ₂ ) × √252",
    "Volatilité annualisée du portefeuille à deux actifs")

# ── 6. EXPORT EXCEL ───────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "7. Export Excel et formules", 1)
add_para(doc,
    "Un export Excel complet a été développé, disponible à la fois depuis l'interface web "
    "(bouton « Excel ») et via un script Python dédié. Le fichier généré contient 4 feuilles "
    "avec des formules Excel natives :"
)

t7 = doc.add_table(rows=1, cols=3)
t7.style = "Table Grid"
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(t7, ["Feuille", "Formules clés", "Description"], [4, 8, 6])
sheets = [
    ("Résumé",       "='VaR 95% 1J'!J8  (liens croisés)",
     "Synthèse des paramètres et résultats, liée aux autres feuilles"),
    ("VaR 95% 1J",   "=SMALL($F$4:$F$5118, J6)\n=AVERAGE($H$4:INDEX(H_range, J6−1))\n=FLOOR(J4*J5, 1)\n=0,5×D+0,5×E",
     "Données, rendements filtrés, colonne triée, VaR/CVaR"),
    ("Performance",  "=C{i-1}*(1+B{i})\n=C{i}/$C$3−1",
     "Valeur cumulée chaînée, rendements cumulés"),
    ("Corrélation",  "(valeurs calculées)",
     "Matrices corrélation/covariance, formule σₚ"),
]
for row_data in sheets:
    add_table_row(t7, row_data)
doc.add_paragraph()

add_para(doc,
    "La colonne H « Rdts triés » utilise la formule =SMALL($F$4:$F$N, ROW()-3), qui génère "
    "automatiquement le tri des rendements sans copier-coller : chaque cellule Hᵢ contient "
    "le i-ème pire rendement. Le k-ième (en rouge) correspond exactement à la VaR, et les "
    "k−1 précédents (en rose) constituent la zone CVaR."
)

# ── 7. CONCLUSION ─────────────────────────────────────────────────────────────
add_heading(doc, "8. Conclusion", 1)
add_para(doc,
    "Ce projet nous a permis d'implémenter de bout en bout un outil professionnel de calcul "
    "de la Value at Risk historique, en partant des formules théoriques du cours jusqu'à une "
    "application web déployable. Les principaux apports techniques sont :"
)
for item in [
    "Implémentation rigoureuse de la méthode SMALL (k = ⌊n × q⌋) conforme au cours Stat M2",
    "Double fenêtre temporelle (VaR 20 ans / perf utilisateur) pour séparer risque et performance",
    "Récupération robuste des données Yahoo Finance via curl_cffi (bypass anti-bot)",
    "Export Excel avec formules natives permettant la vérification et l'audit des calculs",
    "Dashboard analytique avec classement des portefeuilles par niveau de risque",
    "Valorisation par nombre d'actions pour refléter une détention réelle",
]:
    add_bullet(doc, item)

add_para(doc,
    "L'application constitue un outil pédagogique permettant de vérifier, en temps réel et "
    "sur n'importe quel portefeuille d'actions mondiales, les calculs réalisés manuellement "
    "en cours. Elle illustre également les limites de la VaR historique : dépendance à la "
    "période d'estimation, absence de prise en compte de la volatilité conditionnelle, et "
    "caractère non-sous-additif dans certains cas — ouvrant la voie à des extensions telles "
    "que la VaR paramétrique (normale multivariée) ou la VaR Monte Carlo."
)

# ── Pied de page ──────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sidy Laye Sarr  ·  Moulaye Koutam  ·  Abdoulaye Diop  —  M2 Finance EFREI  —  2025/2026")
run.font.size = Pt(9); run.font.color.rgb = LGRAY; run.italic = True

# ── Sauvegarde ────────────────────────────────────────────────────────────────
output = os.path.join(os.path.expanduser("~"), "Documents", "Rapport_VaR_Sarr_Koutam_Diop_v2.docx")
doc.save(output)
print(f"Rapport sauvegardé : {output}")
